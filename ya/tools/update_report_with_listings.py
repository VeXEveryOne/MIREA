from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


FIG_RE = re.compile(r"^Рисунок\s+(\d+)\s*[–-]\s*(.+)$")


@dataclass(frozen=True)
class FigureItem:
    image_blob: bytes
    caption: str


def para_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def para_has_drawing(p: Paragraph) -> bool:
    return bool(p._element.xpath(".//w:drawing"))  # noqa: SLF001


def image_blob_from_paragraph(doc: Document, p: Paragraph) -> bytes | None:
    blips = p._element.xpath('.//*[local-name()="blip"]')  # noqa: SLF001
    if not blips:
        return None
    r_id = blips[0].get(qn("r:embed"))
    if not r_id:
        return None
    part = doc.part.related_parts.get(r_id)
    if part is None:
        return None
    return part.blob


def remove_paragraph(p: Paragraph) -> None:
    el = p._element  # noqa: SLF001
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def set_table_borders_none(table: Table) -> None:
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def insert_table_before(anchor: Paragraph, table: Table) -> None:
    tbl_el = table._tbl  # noqa: SLF001
    tbl_el.getparent().remove(tbl_el)
    anchor._p.addprevious(tbl_el)  # noqa: SLF001


def build_picture_table(
    doc: Document,
    items: list[FigureItem],
    *,
    cols: int,
    picture_width_cm: float,
) -> Table:
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    set_table_borders_none(table)

    for idx, item in enumerate(items):
        r = idx // cols
        c = idx % cols
        cell = table.cell(r, c)
        cell.paragraphs[0].text = ""

        pic_p = cell.paragraphs[0]
        pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic_run = pic_p.add_run()
        pic_run.add_picture(BytesIO(item.image_blob), width=Cm(picture_width_cm))

        cap_p = cell.add_paragraph(item.caption)
        cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap_p.paragraph_format.space_before = 0
        cap_p.paragraph_format.space_after = 0

    return table


def find_figure_caption_paragraphs(doc: Document) -> dict[int, Paragraph]:
    mapping: dict[int, Paragraph] = {}
    for p in doc.paragraphs:
        t = para_text(p)
        m = FIG_RE.match(t)
        if not m:
            continue
        mapping[int(m.group(1))] = p
    return mapping


def find_near_image_paragraph(doc: Document, caption_p: Paragraph) -> Paragraph | None:
    if para_has_drawing(caption_p):
        return caption_p
    # Search a few paragraphs backwards, typical pattern: [picture] then [caption]
    caption_el = caption_p._element  # noqa: SLF001
    paras = list(doc.paragraphs)
    idx = None
    for i, p in enumerate(paras):
        if p._element is caption_el:  # noqa: SLF001
            idx = i
            break
    if idx is None:
        return None
    for back in range(1, 6):
        j = idx - back
        if j < 0:
            break
        p = paras[j]
        if para_has_drawing(p):
            return p
        if para_text(p) != "":
            # Stop at meaningful text.
            break
    # As fallback, look forward a bit.
    for fwd in range(1, 4):
        j = idx + fwd
        if j >= len(paras):
            break
        p = paras[j]
        if para_has_drawing(p):
            return p
        if para_text(p) != "":
            break
    return None


def compact_figures(
    doc: Document,
    *,
    figure_numbers: list[int],
    cols: int,
    picture_width_cm: float,
) -> None:
    captions = find_figure_caption_paragraphs(doc)
    missing = [n for n in figure_numbers if n not in captions]
    if missing:
        raise RuntimeError(f"Missing figure captions: {missing}")

    items: list[FigureItem] = []
    to_remove: list[Paragraph] = []

    for n in figure_numbers:
        cap_p = captions[n]
        img_p = find_near_image_paragraph(doc, cap_p)
        if img_p is None:
            raise RuntimeError(f"Image paragraph not found for figure {n}")
        blob = image_blob_from_paragraph(doc, img_p)
        if blob is None:
            raise RuntimeError(f"Image blob not found for figure {n}")
        items.append(FigureItem(image_blob=blob, caption=para_text(cap_p)))
        if img_p is not cap_p:
            to_remove.append(img_p)
        to_remove.append(cap_p)

    # Insert table before the first caption paragraph.
    anchor = captions[figure_numbers[0]]
    table = build_picture_table(doc, items, cols=cols, picture_width_cm=picture_width_cm)
    insert_table_before(anchor, table)

    # Remove original picture/caption paragraphs.
    # Remove in reverse document order by their position (by element) to avoid shifting issues.
    pos: dict[object, int] = {}
    for i, p in enumerate(doc.paragraphs):
        pos[p._element] = i  # noqa: SLF001

    unique_by_el = {p._element: p for p in to_remove}  # noqa: SLF001
    ordered = sorted(unique_by_el.values(), key=lambda p: pos.get(p._element, -1), reverse=True)  # noqa: SLF001
    for p in ordered:
        remove_paragraph(p)


def add_code_listing(
    doc: Document,
    *,
    anchor: Paragraph,
    caption: str,
    code: str,
    font_size_pt: float = 8.0,
) -> Paragraph:
    cap_p = anchor.insert_paragraph_before(caption, style="Normal")
    cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap_p.paragraph_format.space_before = 0
    cap_p.paragraph_format.space_after = 0

    table = doc.add_table(rows=1, cols=1)
    set_table_borders_none(table)
    cell = table.cell(0, 0)
    cell.text = code.rstrip("\n")

    # Format code paragraphs.
    for p in cell.paragraphs:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(font_size_pt)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")  # noqa: SLF001

    # Move table right after caption.
    tbl_el = table._tbl  # noqa: SLF001
    tbl_el.getparent().remove(tbl_el)
    cap_p._p.addnext(tbl_el)  # noqa: SLF001

    # Return new anchor as paragraph after the table (insert before current anchor each time).
    return cap_p


def read_lines(path: Path, start_line: int, end_line: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    start = max(1, start_line)
    end = min(end_line, len(lines))
    snippet = lines[start - 1 : end]
    return "\n".join(snippet) + "\n"


def find_paragraph(doc: Document, *, startswith: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if para_text(p).startswith(startswith):
            return p
    return None


def find_next_heading1(doc: Document, after: Paragraph) -> Paragraph | None:
    after_el = after._element  # noqa: SLF001
    seen = False
    for p in doc.paragraphs:
        if not seen:
            if p._element is after_el:  # noqa: SLF001
                seen = True
            continue
        if p.style and p.style.name == "Heading 1":
            return p
    return None


def update_report(input_docx: Path, output_docx: Path, code_root: Path) -> None:
    doc = Document(str(input_docx))

    # Compact tall smartphone screenshots to reduce page count.
    compact_figures(doc, figure_numbers=[1, 2, 3], cols=3, picture_width_cm=5.0)
    compact_figures(doc, figure_numbers=[4, 5, 6], cols=3, picture_width_cm=5.0)
    compact_figures(doc, figure_numbers=[7, 8, 9], cols=3, picture_width_cm=5.0)
    compact_figures(doc, figure_numbers=[10, 11, 12], cols=3, picture_width_cm=5.0)
    compact_figures(doc, figure_numbers=list(range(19, 31)), cols=3, picture_width_cm=5.0)

    # Insert code listings into section 3.3 before section 4.
    sec33 = find_paragraph(doc, startswith="3.3 ")
    if sec33 is None:
        raise RuntimeError("Section 3.3 not found")
    next_h1 = find_next_heading1(doc, sec33)
    if next_h1 is None:
        raise RuntimeError("Next Heading 1 after 3.3 not found")

    intro = next_h1.insert_paragraph_before(
        "Ниже приведены листинги (фрагменты) основных компонентов приложения.",
        style="Normal",
    )
    intro.paragraph_format.space_before = 0
    intro.paragraph_format.space_after = 0

    # Listing sources
    main_activity = code_root / "com" / "example" / "playlist_maker_android_albahtinilya" / "MainActivity.kt"
    search_vm = (
        code_root
        / "com"
        / "example"
        / "playlist_maker_android_albahtinilya"
        / "ui"
        / "search"
        / "SearchViewModel.kt"
    )
    tracks_repo = (
        code_root
        / "com"
        / "example"
        / "playlist_maker_android_albahtinilya"
        / "data"
        / "repository"
        / "TracksRepositoryImpl.kt"
    )
    app_db = code_root / "com" / "example" / "playlist_maker_android_albahtinilya" / "data" / "db" / "AppDatabase.kt"
    track_entity = (
        code_root
        / "com"
        / "example"
        / "playlist_maker_android_albahtinilya"
        / "data"
        / "db"
        / "TrackEntity.kt"
    )
    playlist_entity = (
        code_root
        / "com"
        / "example"
        / "playlist_maker_android_albahtinilya"
        / "data"
        / "db"
        / "PlaylistEntity.kt"
    )

    listings = [
        (
            "Листинг 1 – MainActivity.kt (фрагмент: маршрутизация и NavHost)",
            read_lines(main_activity, 140, 245),
        ),
        (
            "Листинг 2 – SearchViewModel.kt (фрагмент: состояние и поиск треков)",
            read_lines(search_vm, 19, 74),
        ),
        (
            "Листинг 3 – TracksRepositoryImpl.kt (фрагмент: поиск и маппинг DTO → Track)",
            read_lines(tracks_repo, 19, 103),
        ),
        ("Листинг 4 – AppDatabase.kt", read_lines(app_db, 1, 200)),
        ("Листинг 5 – TrackEntity.kt", read_lines(track_entity, 1, 200)),
        ("Листинг 6 – PlaylistEntity.kt", read_lines(playlist_entity, 1, 200)),
    ]

    anchor = next_h1
    for caption, code in listings:
        add_code_listing(doc, anchor=anchor, caption=caption, code=code, font_size_pt=8.0)

    doc.save(str(output_docx))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", help="Path to source report .docx")
    parser.add_argument("output", help="Path to output report .docx")
    parser.add_argument(
        "--code-root",
        required=True,
        help="Path to app/src/main/java root (contains com/...)",
    )
    args = parser.parse_args()

    input_docx = Path(args.input)
    output_docx = Path(args.output)
    code_root = Path(args.code_root)

    update_report(input_docx, output_docx, code_root)


if __name__ == "__main__":
    main()
