from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, Iterator, Literal

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm


BlockKind = Literal["p", "t"]


@dataclass(frozen=True)
class Block:
    kind: BlockKind
    obj: Paragraph | Table

    @property
    def element(self):
        return self.obj._element  # noqa: SLF001 (python-docx internal)


W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def iter_blocks(doc: Document) -> Iterator[Block]:
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Block("p", Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield Block("t", Table(child, doc))


def remove_block(block: Block) -> None:
    el = block.element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def para_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def para_has_drawing(p: Paragraph) -> bool:
    # Pictures are stored as w:drawing in paragraph runs.
    # python-docx' xpath helper already knows the 'w' namespace prefix.
    return bool(p._element.xpath(".//w:drawing"))  # noqa: SLF001


def first_cell_text(t: Table) -> str:
    try:
        cell = t.cell(0, 0)
    except Exception:
        return ""
    return " ".join((pp.text or "").strip() for pp in cell.paragraphs).strip()


def find_block_index(
    blocks: list[Block],
    *,
    kind: BlockKind,
    predicate,
) -> int | None:
    for i, b in enumerate(blocks):
        if b.kind != kind:
            continue
        if predicate(b.obj):
            return i
    return None


def delete_range(blocks: list[Block], start: int, end_exclusive: int) -> None:
    for b in blocks[start:end_exclusive]:
        remove_block(b)


def insert_paragraph_after(paragraph: Paragraph, text: str, *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # noqa: SLF001
    new_para = Paragraph(new_p, paragraph._parent)  # noqa: SLF001
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_table_borders_none(table: Table) -> None:
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def image_blob_from_paragraph(doc: Document, p: Paragraph) -> bytes | None:
    # Extract the first embedded image (a:blip r:embed="rIdX") from the paragraph.
    blips = p._element.xpath('.//*[local-name()="blip"]')  # noqa: SLF001
    if not blips:
        return None
    r_id = blips[0].get(qn("r:embed"))
    if not r_id:
        return None
    part = doc.part.related_parts.get(r_id)
    if part is None:
        return None
    return part.blob


def renumber_figure_captions(doc: Document) -> int:
    # Renumber all "Рисунок N – ..." in document order starting from 1.
    pattern = re.compile(r"^Рисунок\s+(\d+)\s*[–-]\s*(.+)$")
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        m = pattern.match(t)
        if not m:
            continue
        n += 1
        rest = m.group(2)
        p.text = f"Рисунок {n} – {rest}"
    return n


def shorten_report(input_path: str, output_path: str) -> None:
    doc = Document(input_path)

    # Rebuild blocks list when needed because doc structure changes.
    blocks = list(iter_blocks(doc))

    # 1) Remove TOC + terms + abbreviations (everything from "СОДЕРЖАНИЕ" to before "ВВЕДЕНИЕ").
    toc_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p) == "СОДЕРЖАНИЕ",
    )
    intro_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p) == "ВВЕДЕНИЕ",
    )
    if toc_i is not None and intro_i is not None and toc_i < intro_i:
        delete_range(blocks, toc_i, intro_i)

    # Remove duplicate technology-choice section (we will add it later as section 4.1).
    blocks = list(iter_blocks(doc))
    sec23_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p).startswith("2.3 ")
        and "программного обеспечения" in para_text(p).lower(),
    )
    sec24_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p).startswith("2.4 ")
        and "Архитектура" in para_text(p),
    )
    if sec23_i is not None and sec24_i is not None and sec23_i < sec24_i:
        delete_range(blocks, sec23_i, sec24_i)

    # 2) Remove the huge stakeholder/persona table.
    blocks = list(iter_blocks(doc))
    for b in blocks:
        if b.kind != "t":
            continue
        if first_cell_text(b.obj).startswith("Стейкхолдер"):
            remove_block(b)
            break

    # 3) Remove analog screenshots (Рисунок 1–12) with their images.
    blocks = list(iter_blocks(doc))
    to_delete: set[int] = set()
    fig_caption = re.compile(r"^Рисунок\s+(\d+)\s*[–-]\s*")
    for i, b in enumerate(blocks):
        if b.kind != "p":
            continue
        t = para_text(b.obj)
        m = fig_caption.match(t)
        if not m:
            continue
        num = int(m.group(1))
        if num > 12:
            continue

        # Delete caption itself.
        to_delete.add(i)

        # Delete preceding empty paragraphs and the paragraph that contains the picture.
        for j in range(i - 1, max(-1, i - 5), -1):
            if j < 0:
                break
            prev = blocks[j]
            if prev.kind != "p":
                break
            prev_p = prev.obj
            if para_has_drawing(prev_p):
                to_delete.add(j)
                break
            if para_text(prev_p) == "":
                to_delete.add(j)
                continue
            break

    for i in sorted(to_delete, reverse=True):
        remove_block(blocks[i])

    # 4) Renumber remaining figures starting from 1.
    total_figures = renumber_figure_captions(doc)

    # 5) Repurpose section 3.2: remove test text, keep screenshots as "Внешний вид приложения".
    blocks = list(iter_blocks(doc))
    sec32_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p).startswith("3.2 ")
        and "Тестирование" in para_text(p),
    )
    if sec32_i is not None and blocks[sec32_i].kind == "p":
        sec32_p: Paragraph = blocks[sec32_i].obj  # type: ignore[assignment]
        sec32_p.text = "3.2 Внешний вид приложения"

        # Find first screenshot block (a paragraph containing drawing or a caption "Рисунок ...")
        start_keep = None
        for j in range(sec32_i + 1, len(blocks)):
            bj = blocks[j]
            if bj.kind != "p":
                continue
            pj: Paragraph = bj.obj  # type: ignore[assignment]
            tj = para_text(pj)
            if para_has_drawing(pj) or tj.startswith("Рисунок "):
                start_keep = j
                break

        if start_keep is not None and start_keep > sec32_i + 1:
            delete_range(blocks, sec32_i + 1, start_keep)

        # Insert short intro text right after the heading.
        intro = insert_paragraph_after(
            sec32_p,
            "Далее представлен внешний вид приложения (скриншоты основных экранов).",
            style="Normal",
        )
        # Keep minimal spacing.
        intro.paragraph_format.space_before = 0
        intro.paragraph_format.space_after = 0

    # 5.1) Compact screenshots: place 2 screenshots per row in a borderless table.
    # This is the main lever to reduce the page count under the 30-page limit.
    blocks = list(iter_blocks(doc))
    sec32_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p) == "3.2 Внешний вид приложения",
    )
    sec33_i = find_block_index(
        blocks,
        kind="p",
        predicate=lambda p: para_text(p).startswith("3.3 ") and "листинг" in para_text(p).lower(),
    )
    if sec32_i is not None and sec33_i is not None and sec32_i < sec33_i:
        screenshots: list[tuple[bytes, str]] = []
        del_idx: set[int] = set()

        i = sec32_i + 1
        while i < sec33_i:
            b = blocks[i]
            if b.kind != "p":
                i += 1
                continue
            p: Paragraph = b.obj  # type: ignore[assignment]
            if not para_has_drawing(p):
                i += 1
                continue

            blob = image_blob_from_paragraph(doc, p)
            if blob is None:
                i += 1
                continue

            # Find the next caption paragraph "Рисунок N – ...".
            caption_i = None
            for j in range(i + 1, min(sec33_i, i + 6)):
                bj = blocks[j]
                if bj.kind != "p":
                    continue
                tj = para_text(bj.obj)
                if tj.startswith("Рисунок "):
                    caption_i = j
                    break

            if caption_i is None:
                i += 1
                continue

            caption_text = para_text(blocks[caption_i].obj)
            screenshots.append((blob, caption_text))

            # Mark image paragraph + any empty paragraphs up to caption + caption itself for deletion.
            for j in range(i, caption_i + 1):
                if blocks[j].kind != "p":
                    continue
                pj: Paragraph = blocks[j].obj  # type: ignore[assignment]
                if para_has_drawing(pj) or para_text(pj) == "" or para_text(pj).startswith("Рисунок "):
                    del_idx.add(j)

            i = caption_i + 1

        # Delete old screenshot paragraphs (images + captions).
        for j in sorted(del_idx, reverse=True):
            remove_block(blocks[j])

        # Insert a 2-column table after the intro paragraph.
        if screenshots:
            # Find intro paragraph (inserted earlier).
            intro_p = None
            for p in doc.paragraphs:
                if (p.text or "").strip().startswith("Далее представлен внешний вид приложения"):
                    intro_p = p
                    break
            if intro_p is not None:
                rows = (len(screenshots) + 1) // 2
                table = doc.add_table(rows=rows, cols=2)
                set_table_borders_none(table)

                for idx, (blob, cap) in enumerate(screenshots):
                    r = idx // 2
                    c = idx % 2
                    cell = table.cell(r, c)
                    # Clear default empty paragraph text.
                    cell.paragraphs[0].text = ""
                    pic_p = cell.paragraphs[0]
                    pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = pic_p.add_run()
                    run.add_picture(BytesIO(blob), width=Cm(7.5))

                    cap_p = cell.add_paragraph(cap)
                    cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cap_p.paragraph_format.space_before = 0
                    cap_p.paragraph_format.space_after = 0

                # Move the table element right after the intro paragraph.
                tbl_el = table._tbl  # noqa: SLF001
                tbl_el.getparent().remove(tbl_el)
                intro_p._p.addnext(tbl_el)  # noqa: SLF001

    # 6) Rename section 3.3 and remove code tables/listing captions (leave only file list & structure).
    doc_blocks = list(iter_blocks(doc))
    sec33_i = find_block_index(
        doc_blocks,
        kind="p",
        predicate=lambda p: para_text(p).startswith("3.3 ")
        and "листинг" in para_text(p).lower(),
    )
    # In some variants heading can be "3.3 Полный листинг ..."
    if sec33_i is None:
        sec33_i = find_block_index(
            doc_blocks,
            kind="p",
            predicate=lambda p: para_text(p).startswith("3.3 ")
            and "листинг" in para_text(p).lower(),
        )

    if sec33_i is not None and doc_blocks[sec33_i].kind == "p":
        sec33_p: Paragraph = doc_blocks[sec33_i].obj  # type: ignore[assignment]
        sec33_p.text = "3.3 Листинг приложения"

        # Determine end of this section: before next Heading 1 (ЗАКЛЮЧЕНИЕ) or end.
        end_i = None
        for j in range(sec33_i + 1, len(doc_blocks)):
            bj = doc_blocks[j]
            if bj.kind != "p":
                continue
            pj: Paragraph = bj.obj  # type: ignore[assignment]
            if (pj.style and pj.style.name == "Heading 1") and para_text(pj) in {
                "ЗАКЛЮЧЕНИЕ",
                "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            }:
                end_i = j
                break
        if end_i is None:
            end_i = len(doc_blocks)

        # Delete listing captions and 1x1 code tables inside this section.
        # Also delete the paragraph mentioning local path to repo.
        to_del: list[Block] = []
        for b in doc_blocks[sec33_i + 1 : end_i]:
            if b.kind == "p":
                t = para_text(b.obj)
                if t.startswith("Листинг "):
                    to_del.append(b)
                if "D:\\GitHub" in t or "D:/GitHub" in t:
                    to_del.append(b)
            elif b.kind == "t":
                if len(b.obj.rows) == 1 and len(b.obj.columns) == 1:
                    to_del.append(b)

        for b in to_del:
            remove_block(b)

        # Insert a compact file list (like in the example report).
        file_lines = [
            "В этом разделе перечислены основные файлы проекта и их назначение:",
            "MainActivity.kt — точка входа в приложение; настройка темы и графа навигации.",
            "ui/search/SearchScreen (в проекте реализован как Composable) — экран поиска треков.",
            "ui/search/DetailScreen.kt — экран карточки трека и действия «в избранное/в плейлист».",
            "ui/playlists/* — экраны списка плейлистов, деталей плейлиста, создание плейлиста.",
            "data/network/ITunesApiService.kt — интерфейс сетевого API (Retrofit).",
            "data/db/AppDatabase.kt — конфигурация локальной БД Room.",
            "data/db/TrackEntity.kt, PlaylistEntity.kt — сущности таблиц БД.",
            "data/preferences/* — хранение истории поиска и настроек (DataStore).",
        ]

        anchor = sec33_p
        for line in file_lines:
            style = "List Bullet" if "—" in line and not line.endswith(":") else "Normal"
            anchor = insert_paragraph_after(anchor, line, style=style)

    # 7) Insert sections 4 and 5 before conclusion (to match the example structure).
    conclusion = None
    for p in doc.paragraphs:
        if (p.style and p.style.name == "Heading 1") and para_text(p) == "ЗАКЛЮЧЕНИЕ":
            conclusion = p
            break

    if conclusion is not None:
        insert_items: list[tuple[str, str]] = [
            ("Heading 1", "4. РЕАЛИЗАЦИЯ МОБИЛЬНОГО ПРИЛОЖЕНИЯ"),
            ("Heading 2", "4.1 Выбор языков программирования и технологий"),
            (
                "Normal",
                "Для реализации приложения выбран язык Kotlin и современный декларативный UI-фреймворк Jetpack Compose. Для асинхронной логики применяются Coroutines и Flow, для навигации — Navigation Compose.",
            ),
            ("Heading 2", "4.2 Реализация визуальных элементов"),
            (
                "Normal",
                "Экранные формы реализованы как Composable-функции с использованием компонентов Material 3. Состояния (загрузка/ошибка/контент) отображаются реактивно в зависимости от данных из ViewModel.",
            ),
            ("Heading 2", "4.3 Реализация сервисов и бизнес-логики"),
            (
                "Normal",
                "Сетевой поиск реализован через Retrofit/OkHttp (iTunes Search API), бизнес-операции инкапсулированы в репозиториях и ViewModel. Для действий «в избранное» и «в плейлист» используются единые модели данных трека и плейлиста.",
            ),
            ("Heading 2", "4.4 Хранение и обработка данных"),
            (
                "Normal",
                "Локальные данные (избранное, плейлисты и их содержимое) сохраняются в Room. История поиска и настройки темы хранятся в DataStore Preferences, что обеспечивает сохранность данных между запусками приложения.",
            ),
            ("Heading 1", "5. ТЕСТИРОВАНИЕ И ОТЛАДКА ПРОГРАММНОГО ПРОДУКТА"),
            (
                "Normal",
                "Тестирование проводилось на эмуляторе Android и на реальном устройстве. Проверялись основные пользовательские сценарии:",
            ),
            ("List Bullet", "поиск треков (успешный поиск, пустой результат, ошибка сети)"),
            ("List Bullet", "переход на карточку трека и корректность отображения полей"),
            ("List Bullet", "добавление/удаление трека в «Избранное» и обновление списка"),
            ("List Bullet", "создание плейлиста и добавление/удаление треков в плейлист"),
            ("List Bullet", "настройки: смена темы, шаринг, поддержка, соглашение"),
            (
                "Normal",
                "По итогам тестирования приложение корректно обрабатывает ошибки сети, сохраняет данные между запусками и обеспечивает стабильную навигацию между экранами.",
            ),
        ]

        for style, text in insert_items:
            conclusion.insert_paragraph_before(text, style=style)

    # 8) Bring main part headings closer to the example formatting.
    for p in doc.paragraphs:
        if not (p.style and p.style.name == "Heading 1"):
            continue
        t = para_text(p)
        if t == "1 ИССЛЕДОВАТЕЛЬСКАЯ ЧАСТЬ":
            p.text = "1. ИССЛЕДОВАТЕЛЬСКАЯ ЧАСТЬ"
        elif t == "3 ТЕХНОЛОГИЧЕСКАЯ ЧАСТЬ":
            p.text = "3. ПРОЕКТИРОВАНИЕ МОБИЛЬНОГО ПРОГРАММНОГО КОМПЛЕКСА"

    for p in doc.paragraphs:
        if not (p.style and p.style.name == "Heading 2"):
            continue
        t = para_text(p)
        if t.startswith("2.4 ") and "Архитектура" in t:
            p.text = t.replace("2.4 ", "2.3 ", 1)

    # Add a short reference to detailed IDEF0 descriptions (external txt file).
    for i, p in enumerate(doc.paragraphs):
        if (p.style and p.style.name == "Heading 2") and para_text(p).startswith("2.1 ") and "IDEF0" in para_text(p):
            # Insert at end of IDEF0 section: just before next Heading 2 or end.
            insert_before = None
            for j in range(i + 1, len(doc.paragraphs)):
                pj = doc.paragraphs[j]
                if pj.style and pj.style.name.startswith("Heading"):
                    insert_before = pj
                    break
            if insert_before is None:
                insert_before = doc.paragraphs[-1]
            insert_before.insert_paragraph_before(
                "Подробное текстовое описание декомпозиций A1–A4 приведено в файле idef0_текстовое_описание.txt.",
                style="Normal",
            )
            break

    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"Figures (after cleanup): {total_figures}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", help="Path to source .docx")
    parser.add_argument("output", help="Path to output .docx")
    args = parser.parse_args()
    shorten_report(args.input, args.output)


if __name__ == "__main__":
    main()
